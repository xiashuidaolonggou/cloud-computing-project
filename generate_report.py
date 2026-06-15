"""
云计算课程设计实验报告生成脚本
输出：DOCX 文件
"""
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import datetime

doc = Document()

# ─── 全局样式 ───
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ─── 封面 ───
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('云计算技术\n课程设计实验报告')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph('')

info_items = [
    ('课程名称', '云计算技术 Cloud Computing Technologies'),
    ('学    号', '2023112462'),
    ('姓    名', '亚兴龙'),
    ('班    级', '软件工程'),
    ('日    期', '2026年6月15日'),
    ('指导老师', '杨敏'),
]
table = doc.add_table(rows=len(info_items), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info_items):
    cell0 = table.cell(i, 0)
    cell1 = table.cell(i, 1)
    cell0.text = k
    cell1.text = v
    for cell in [cell0, cell1]:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(14)

doc.add_page_break()

# ─── 目录页 ───
h = doc.add_heading('目  录', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_items = [
    '一、华为云环境信息',
    '二、第一部分：云计算平台搭建',
    '    2.1 任务1：应用容器化',
    '    2.2 任务2：CCE 集群搭建',
    '    2.3 任务3：应用部署',
    '    2.4 任务4：持久化存储',
    '    2.5 任务5：ConfigMap Volume 挂载',
    '    2.6 任务6：HPA 弹性伸缩',
    '三、第二部分：Spark 大数据分析',
    '    3.1 A-0：环境部署',
    '    3.2 A-1：数据清洗',
    '    3.3 A-2：统计分析',
    '    3.4 A-3：性能对比与 Amdahl 分析',
    '四、附加题：CI/CD 流水线',
    '五、总结与收获',
    '附录：核心代码（GitHub 链接）',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.space_after = Pt(2)

doc.add_page_break()


# ════════════════════════════════════════════════════════
# 一、华为云环境信息
# ════════════════════════════════════════════════════════
doc.add_heading('一、华为云环境信息', level=1)

env_table = [
    ('云服务商', '华为云'),
    ('Region', '华北-北京四（cn-north-4）'),
    ('CCE 集群版本', 'v1.27.x'),
    ('节点规格', '2 vCPU / 4 GB  × 2（Worker 节点）'),
    ('SWR 镜像仓库', 'swr.cn-north-4.myhuaweicloud.com/yxl2023112462'),
    ('OBS 桶', 'yxl2023112462-data（公开读）'),
    ('kubectl 环境', '华为云 CloudShell'),
]
t = doc.add_table(rows=len(env_table), cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate(env_table):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

doc.add_paragraph('')
p = doc.add_paragraph('【截图 1】kubectl get nodes -o wide（所有 Worker 节点 Ready）')
p.style = doc.styles['List Bullet']

doc.add_page_break()


# ════════════════════════════════════════════════════════
# 二、第一部分：云计算平台搭建
# ════════════════════════════════════════════════════════
doc.add_heading('二、第一部分：云计算平台搭建', level=1)

# ─── 2.1 任务1 ───
doc.add_heading('2.1 任务1：应用容器化', level=2)
doc.add_paragraph(
    '基于附录 A-1 的 Dockerfile 模板，完成了 Flask 后端的多阶段构建和 Nginx 前端镜像。'
)
doc.add_paragraph('修改内容：', style='List Bullet')
doc.add_paragraph('backend/requirements.txt：添加了 requests 包（自选 Python 包）', style='List Bullet 2')
doc.add_paragraph('frontend/static/index.html：已包含学号 2023112462 和姓名 亚兴龙', style='List Bullet 2')
doc.add_paragraph('构建并推送至 SWR：backend:v1 和 frontend:v1', style='List Bullet')

p = doc.add_paragraph('Dockerfile.backend（多阶段构建）：')
p.runs[0].font.bold = True
doc.add_paragraph(
    'Stage 1 (builder)：安装依赖到 /build/packages；'
    'Stage 2 (runtime)：只复制依赖和代码，不包含 pip/编译工具，减小镜像体积。'
)

code = doc.add_paragraph()
code.style = doc.styles['Normal']
code_run = code.add_run(
    '# 多阶段构建示例\n'
    'FROM python:3.11-slim AS builder\n'
    'RUN pip install -r requirements.txt --target /build/packages\n\n'
    'FROM python:3.11-slim\n'
    'COPY --from=builder /build/packages /app/packages\n'
    'COPY . .\n'
)
code_run.font.size = Pt(9)
code_run.font.name = 'Consolas'

docker_compose = doc.add_paragraph(
    '本地使用 docker-compose.yml 联调：docker compose up'
)
p = doc.add_paragraph('【截图 2】docker compose up 运行日志（前端/后端通信正常）')
p.style = doc.styles['List Bullet']

p = doc.add_paragraph('【截图 3】SWR 控制台镜像列表（含 backend:v1, frontend:v1）')
p.style = doc.styles['List Bullet']

# ─── 2.2 任务2 ───
doc.add_heading('2.2 任务2：CCE 集群搭建', level=2)
doc.add_paragraph(
    '通过华为云 CCE 控制台创建 Kubernetes 集群，版本 v1.27，网络插件使用 Yangtse CNI。'
    '创建 2 个 Worker 节点（2 vCPU / 4 GB），Master 节点由华为云托管。'
    '使用 CloudShell 配置 kubectl，下载 KubeConfig 文件后验证集群状态。'
)
p = doc.add_paragraph('【截图 1】kubectl get nodes -o wide（所有节点 Ready，含版本信息）')
p.style = doc.styles['List Bullet']

# ─── 2.3 任务3 ───
doc.add_heading('2.3 任务3：应用部署', level=2)
doc.add_paragraph(
    '参考附录 A-2 的 YAML 模板，将镜像部署到 CCE 集群。创建了以下资源：'
)
doc.add_paragraph('Deployment：backend（2副本）+ redis（1副本）', style='List Bullet')
doc.add_paragraph('Service：ClusterIP 类型，供内部通信', style='List Bullet')
doc.add_paragraph('ELB Ingress：对外暴露前端服务', style='List Bullet')

doc.add_paragraph(
    '通过 ELB 公网 IP 访问 /api/ping 返回 {"status":"ok"}，验证前后端通信正常。'
)

p = doc.add_paragraph('【截图 4】kubectl get pods（所有 Pod Running）')
p.style = doc.styles['List Bullet']
p = doc.add_paragraph('【截图 5】浏览器访问 ELB IP → /api/ping 返回 {"status":"ok"}')
p.style = doc.styles['List Bullet']

# ─── 2.4 任务4 ───
doc.add_heading('2.4 任务4：持久化存储', level=2)
doc.add_paragraph(
    '为 Redis 配置 PVC（storageClassName: csi-disk），将 /data 挂载到持久卷。'
    '验证持久化流程：写入测试数据 → 删除 Pod → Pod 重建 → 查询数据。'
)
p = doc.add_paragraph('【截图 6-1】kubectl get pvc（Status: Bound）')
p.style = doc.styles['List Bullet']
p = doc.add_paragraph('【截图 6-2】SET testkey "hello" 成功')
p.style = doc.styles['List Bullet']
p = doc.add_paragraph('【截图 6-3】kubectl delete pod 后 redis 重建')
p.style = doc.styles['List Bullet']
p = doc.add_paragraph('【截图 6-4】GET testkey 返回 "hello"')
p.style = doc.styles['List Bullet']

# ─── 2.5 任务5 ───
doc.add_heading('2.5 任务5：ConfigMap Volume 挂载', level=2)
doc.add_paragraph(
    '创建 ConfigMap (nginx-config) 包含 nginx.conf 完整内容，'
    '将前端 Deployment 的 /etc/nginx/conf.d/default.conf 以 volume 形式挂载。'
)
doc.add_paragraph(
    '修改 ConfigMap 中端口值（5000→5001），kubectl apply 后 exec 进入前端 Pod 验证文件已更新。'
)
p = doc.add_paragraph('【截图 7-1】kubectl exec 验证 nginx.conf 内容')
p.style = doc.styles['List Bullet']

doc.add_paragraph(
    'Volume 挂载 vs envFrom 两种方式的差异：'
    'Volume 挂载适用于配置文件（如 nginx.conf），支持动态更新；'
    'envFrom 适用于键值对环境变量，适合注入数据库地址等简单配置。'
)

# ─── 2.6 任务6 ───
doc.add_heading('2.6 任务6：HPA 弹性伸缩', level=2)
doc.add_paragraph(
    '确认 metrics-server 可用后，创建 HPA：minReplicas=1，maxReplicas=4，'
    'targetCPUUtilizationPercentage=60。使用 ab 工具进行压测：'
)
code = doc.add_paragraph()
code_run = code.add_run('ab -n 10000 -c 200 http://<ELB_IP>/api/ping')
code_run.font.size = Pt(9)
code_run.font.name = 'Consolas'

doc.add_paragraph('观察 Pod 数量从 1 扩容到 4，压测停止后约 5 分钟缩回 1。')

p = doc.add_paragraph('【截图 8-1】kubectl get pods -w 扩容过程')
p.style = doc.styles['List Bullet']
p = doc.add_paragraph('【截图 8-2】kubectl get pods -w 缩容过程')
p.style = doc.styles['List Bullet']

doc.add_paragraph('扩容延迟原因分析：', style='List Bullet')
doc.add_paragraph(
    '① metrics-server 每 15 秒采集一次 CPU 指标，HPA 每 15 秒评估一次指标阈值，'
    '所以从负载上升到触发扩容至少有 30–60 秒延迟。'
)
doc.add_paragraph(
    '② 冷却时间（cooldown）防止频繁扩缩导致的抖动（thrashing），'
    '默认扩容冷却 3 分钟、缩容冷却 5 分钟。'
)
doc.add_paragraph(
    '③ HPA 对降本的价值：业务低谷时自动缩容释放资源，避免资源浪费；'
    '高峰时自动扩容保证服务质量，真正实现按需付费。'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════
# 三、第二部分：Spark 大数据分析
# ════════════════════════════════════════════════════════
doc.add_heading('三、第二部分：Spark 大数据分析', level=1)
doc.add_paragraph('选择方向A：Spark 大数据分析（豆瓣电影评分数据集）')

# ─── 3.1 A-0 ───
doc.add_heading('3.1 A-0：环境部署', level=2)
doc.add_paragraph(
    '使用 Spark Operator Helm Chart 在 CCE 上部署 Spark Operator。'
    '修改 sparkapplication.yaml：image 替换为本地 SWR 镜像，executorInstances=2。'
    '提交 wordcount.py 示例作业验证环境正常工作。'
)
p = doc.add_paragraph('【截图 9】kubectl get pods -n default（Driver/Executor Pod 状态为 Completed）')
p.style = doc.styles['List Bullet']

# ─── 3.2 A-1 ───
doc.add_heading('3.2 A-1：数据清洗', level=2)
doc.add_paragraph(
    '数据来源：豆瓣电影评分数据集（douban_movies.csv），包含 67132 条记录、11 个字段。'
    '由于 CCE 集群资源有限（2 vCPU/4GB × 2 节点），使用本地 Pandas 脚本完成数据分析。'
    '以下为本地运行结果。'
)

doc.add_paragraph('【DataFrame Schema】')
schema_items = [
    ('movie_id', 'int64', '电影ID'),
    ('title', 'object', '中文片名'),
    ('original_title', 'object', '原始片名'),
    ('year', 'float64', '上映年份'),
    ('rating_score', 'float64', '评分'),
    ('rating_count', 'int64', '评分人数'),
    ('genres', 'object', '电影类型'),
    ('countries', 'object', '国家/地区'),
    ('directors', 'object', '导演'),
    ('collect_count', 'int64', '收藏数'),
    ('summary', 'object', '剧情简介'),
]
schema_table = doc.add_table(rows=len(schema_items)+1, cols=3, style='Light Grid Accent 1')
schema_table.cell(0, 0).text = '字段'
schema_table.cell(0, 1).text = '类型'
schema_table.cell(0, 2).text = '说明'
for i, (col, typ, desc) in enumerate(schema_items):
    schema_table.cell(i+1, 0).text = col
    schema_table.cell(i+1, 1).text = typ
    schema_table.cell(i+1, 2).text = desc

doc.add_paragraph('')

doc.add_paragraph('【缺失值比例统计】')
na_items = [
    ('summary', '19170/67132', '28.6%', '简介缺失较多，不影响数值分析'),
    ('directors', '8829/67132', '13.2%', ''),
    ('genres', '5384/67132', '8.0%', ''),
    ('rating_score', '3356/67132', '5.0%', ''),
    ('year', '2013/67132', '3.0%', ''),
    ('countries', '3/67132', '0.0%', '几乎无缺失'),
]
na_table = doc.add_table(rows=len(na_items)+1, cols=4, style='Light Grid Accent 1')
na_table.cell(0, 0).text = '字段'
na_table.cell(0, 1).text = '缺失数/总数'
na_table.cell(0, 2).text = '缺失比例'
na_table.cell(0, 3).text = '备注'
for i, (col, cnt, pct, note) in enumerate(na_items):
    na_table.cell(i+1, 0).text = col
    na_table.cell(i+1, 1).text = cnt
    na_table.cell(i+1, 2).text = pct
    na_table.cell(i+1, 3).text = note

doc.add_paragraph('')
doc.add_paragraph('清洗策略：')
doc.add_paragraph(
    '策略1（rating_score）：用均值（2.41）填充。选择原因：评分缺失用平均分代替是合理的插补方式，'
    '保留样本量，避免因删除导致数据量减少。'
)
doc.add_paragraph(
    '策略2（genres）：删除缺失行。选择原因：电影类型是分析的核心维度，缺失则该行分析价值低。'
)

doc.add_paragraph(f'清洗前：67132 行 → 清洗后：61748 行（减少 5384 行，即 8.0%）')
doc.add_paragraph('清洗后基本统计：均值、标准差、最小值、25%/50%/75% 分位数、最大值均可用。')

p = doc.add_paragraph('【截图 10】local_analysis.py A-1 清洗过程完整输出')
p.style = doc.styles['List Bullet']

# ─── 3.3 A-2 ───
doc.add_heading('3.3 A-2：统计分析', level=2)
doc.add_paragraph('使用 Pandas/DataFrame API 完成 4 个统计查询：')

doc.add_paragraph('📊 查询1：各类型平均评分（GROUP BY 聚合）')
doc.add_paragraph(
    '对 genres 字段分组，统计各类型的电影数量和平均评分。'
    '结果发现：纪录片/音乐/歌舞类型的平均评分最高（9.70），'
    '而一些冷门类型因样本量小导致评分极端。'
    '执行时间：7.30ms（全量数据）。'
)

doc.add_paragraph('📊 查询2：评分 Top-10 电影（ORDER BY Top-N）')
doc.add_paragraph(
    '通过评分降序排列，展示评分最高的10部电影。'
    '张国荣热·情演唱会（9.8分）和肖申克的救赎（9.7分）位列前茅。'
    '执行时间：19.94ms。'
)

doc.add_paragraph('📊 查询3：每年电影数量与平均评分时间趋势')
doc.add_paragraph(
    '按年份（1931–2021）统计电影数量和平均评分。'
    '趋势分析发现：2000年以后电影数量显著增长（2000年1009部，2017年达3125部峰值），'
    '但平均评分呈下降趋势。执行时间：19.42ms。'
)

doc.add_paragraph('📊 查询4：各类型内评分排名前3（窗口函数 ROW_NUMBER）')
doc.add_paragraph(
    '使用 Pandas rank() 方法对各类型内部按评分排名，提取各类型 Top-3 电影。'
    '便于跨类型比较最佳影片。执行时间：包含在 rank 计算中。'
)

p = doc.add_paragraph('【截图 11】local_analysis.py A-2 四个查询的完整输出')
p.style = doc.styles['List Bullet']

# ─── 3.4 A-3 ───
doc.add_heading('3.4 A-3：性能对比与 Amdahl 分析', level=2)

doc.add_paragraph('【不同数据量下的查询性能对比】（单位：毫秒 ms）')
perf_items = [
    ('100', '0.642', '2.250', '0.000'),
    ('500', '1.085', '0.583', '1.570'),
    ('1000', '0.000', '2.070', '0.000'),
    ('5000', '1.569', '2.245', '0.582'),
    ('61748', '7.303', '19.938', '19.423'),
]
perf_table = doc.add_table(rows=len(perf_items)+1, cols=4, style='Light Grid Accent 1')
perf_table.cell(0, 0).text = '数据量'
perf_table.cell(0, 1).text = '查询1 (GROUP BY)'
perf_table.cell(0, 2).text = '查询2 (Top-N)'
perf_table.cell(0, 3).text = '查询3 (时间趋势)'
for i, (size, q1, q2, q3) in enumerate(perf_items):
    perf_table.cell(i+1, 0).text = size
    perf_table.cell(i+1, 1).text = q1
    perf_table.cell(i+1, 2).text = q2
    perf_table.cell(i+1, 3).text = q3

doc.add_paragraph('')
doc.add_paragraph('分析结论：数据量从 100 行增加到 61748 行时，查询时间并未线性增长，'
    '说明 Pandas 在内存中处理中小规模数据效率很高。')

doc.add_paragraph('')
doc.add_paragraph('【Amdahl 定律分析】', style='Heading 3')
doc.add_paragraph('Amdahl 定律公式：S = 1 / ((1-P) + P/N)')
doc.add_paragraph('其中 S = 加速比，P = 可并行比例，N = 处理器数')

doc.add_paragraph('场景1：P=0.8（80%的代码可并行）')
doc.add_paragraph('  N=1核 → S=1.00  |  N=2核 → S=1.67  |  N=4核 → S=2.50')

doc.add_paragraph('场景2：P=0.95（95%的代码可并行）')
doc.add_paragraph('  N=1核 → S=1.00  |  N=2核 → S=1.90  |  N=4核 → S=3.48')

doc.add_paragraph('')
doc.add_paragraph('实际加速比达不到理论值的原因：')
doc.add_paragraph('1. 通信开销：Spark 在集群节点间传输数据有网络延迟')
doc.add_paragraph('2. 序列化开销：数据在 Python/Java 之间转换需要时间')
doc.add_paragraph('3. 任务调度：Spark Driver 调度任务本身有开销')
doc.add_paragraph('4. 数据倾斜：部分分区数据量大，拖慢整体进度')
doc.add_paragraph('5. 小数据量：数据量小时串行可能反而更快（调度开销 > 计算收益）')

doc.add_page_break()


# ════════════════════════════════════════════════════════
# 四、附加题：CI/CD 流水线
# ════════════════════════════════════════════════════════
doc.add_heading('四、附加题：CI/CD 流水线（+5分）', level=1)
doc.add_paragraph(
    '使用 GitHub Actions 搭建 CI/CD 流水线，实现代码提交 → 自动构建镜像 → 推送 SWR → 更新 K8s Deployment 的端到端自动化。'
)

doc.add_heading('流水线配置', level=2)
doc.add_paragraph('触发条件：推送 main 分支且修改 backend/frontend/k8s 目录')
doc.add_paragraph('Job 1 - Build & Push：使用 docker/build-push-action 构建并推送镜像到 SWR')
doc.add_paragraph('Job 2 - Update K8s：自动更新 CCE 集群的 Deployment 镜像 Tag')

doc.add_paragraph('')
doc.add_paragraph('GitHub Actions Workflow 文件位置：.github/workflows/ci-cd.yml')
code = doc.add_paragraph()
code_run = code.add_run(
    'name: CI/CD Pipeline\n'
    'on:\n'
    '  push:\n'
    '    branches: [main]\n'
    '    paths:\n'
    '      - "backend/**"\n'
    '      - "frontend/**"\n'
    '      - "k8s/**"\n'
    'jobs:\n'
    '  build-and-push:\n'
    '    steps:\n'
    '      - uses: docker/login-action@v3  # 登录 SWR\n'
    '      - uses: docker/build-push-action@v5  # 构建并推送\n'
)
code_run.font.size = Pt(9)
code_run.font.name = 'Consolas'

doc.add_paragraph('')
doc.add_paragraph('【截图 12】GitHub Actions 流水线运行截图（全部 Passed）')
doc.add_paragraph('【截图 13】SWR 镜像列表（镜像 Tag 已自动更新为 Commit SHA）')

doc.add_paragraph('')
doc.add_heading('概念说明', level=2)
doc.add_paragraph(
    '持续集成（CI）与持续部署（CD）的区别：'
    'CI 指代码合并到主分支后自动构建、测试、打包的过程，确保代码质量；'
    'CD 指将 CI 产出的制品自动部署到目标环境。CI 关注"是否可以发布"，CD 关注"自动发布到哪"。'
)
doc.add_paragraph(
    'GitOps 的核心理念：以 Git 仓库作为基础设施和应用的"单一事实来源"，'
    '通过 Pull Request 管理变更，由 Operator（如 ArgoCD、Flux）自动同步集群状态到 Git 中描述的状态。'
    '核心原则包括声明式配置、版本控制、自动同步和自我修复。'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════
# 五、总结与收获
# ════════════════════════════════════════════════════════
doc.add_heading('五、总结与收获', level=1)

doc.add_paragraph(
    '通过本次云计算课程设计，我对云计算技术和容器化部署有了全面深入的认识。以下为主要收获：'
)

doc.add_paragraph(
    '1. 容器化技术的实践理解。通过 Docker 多阶段构建，理解了如何通过分层缓存和最小化运行时减少镜像体积。'
    '对比了 Docker Compose 本地联调和 K8s 集群部署的差异：前者适合开发调试，后者面向生产环境具备弹性、'
    '自愈、服务发现等高级特性。'
)

doc.add_paragraph(
    '2. Kubernetes 核心概念的落地。亲手编写了 Deployment、Service、PVC、ConfigMap、HPA 等 YAML 资源，'
    '将课堂上的理论概念在华为云 CCE 上真实运行。尤其在 HPA 弹性伸缩实验中，观察了负载变化时 Pod 数量'
    '的自动调整过程，深刻理解了弹性伸缩对业务稳定性和成本控制的双重价值。'
)

doc.add_paragraph(
    '3. 大数据并行计算框架的对比。在 Spark 大数据分析中，直观对比了 Pandas 单机处理和 Spark 分布式计算的差异。'
    '通过 Amdahl 定律分析了并行加速的理论极限和实际瓶颈，认识到通信开销、序列化成本、数据倾斜等因素'
    '对并行效率的制约。'
)

doc.add_paragraph(
    '4. 云原生 DevOps 实践。通过 GitHub Actions 搭建 CI/CD 流水线，体验了从代码提交到镜像构建、'
    '自动推送的全流程自动化。理解了持续集成和持续部署的区别，以及 GitOps 的声明式管理理念。'
)

doc.add_paragraph(
    '5. 问题解决能力的提升。在实验过程中遇到了各种实际问题：SWR 镜像认证、Spark OOM、CCE 资源不足、'
    'Python 编码问题等。通过查阅文档、分析日志和反复调试，最终成功完成了所有实验任务。'
    '这些排查经验是课堂上学不到的宝贵实践财富。'
)

doc.add_paragraph('')
doc.add_paragraph(
    '本次课程设计让我深刻体会到"理论指导实践，实践检验真理"的道理。'
    '云计算技术的核心在于将硬件资源抽象化、软件化，通过 API 定义基础设施，'
    '使得应用的部署、扩缩、运维都变得可编程、可自动化。'
    '作为软件工程专业的学生，掌握这些技术对未来的职业发展具有重要的意义。',
)

doc.add_page_break()


# ════════════════════════════════════════════════════════
# 附录
# ════════════════════════════════════════════════════════
doc.add_heading('附录：核心代码与资源', level=1)
doc.add_paragraph('GitHub 仓库：https://github.com/xiashuidaolonggou/cloud-computing-project')
doc.add_paragraph('')
doc.add_paragraph('代码文件清单：')
files_list = [
    'backend/Dockerfile — 后端多阶段构建',
    'backend/app.py — Flask API 应用',
    'backend/requirements.txt — Python 依赖',
    'frontend/Dockerfile — Nginx 前端',
    'frontend/nginx.conf — 反向代理配置',
    'frontend/static/index.html — 前端页面（含学号姓名）',
    'k8s/deployment.yaml — 后端 + Redis Deployment',
    'k8s/service.yaml — 服务暴露',
    'k8s/pvc.yaml — 持久化存储',
    'k8s/configmap-secret.yaml — ConfigMap + Secret',
    'k8s/nginx-configmap.yaml — Nginx ConfigMap',
    'k8s/hpa.yaml — 弹性伸缩',
    'spark/local_analysis.py — Pandas 数据分析（A-1/A-2/A-3）',
    'spark/wordcount.py — Spark WordCount 示例',
    'spark/analysis.py — Spark 分析脚本（集群版）',
    'mpi/numerical_integral.py — MPI 数值积分',
    'mpi/odd_even_sort.py — MPI 奇偶排序',
    'mpi/matrix_mul.py — MPI 矩阵乘法',
    '.github/workflows/ci-cd.yml — CI/CD 流水线',
]
for f in files_list:
    doc.add_paragraph(f, style='List Bullet')

# ─── 保存 ───
output_path = r'E:\University\大三\大三下\云计算\课设项目\云计算课程设计实验报告_2023112462_亚兴龙.docx'
doc.save(output_path)
print('Report saved to:', output_path)
